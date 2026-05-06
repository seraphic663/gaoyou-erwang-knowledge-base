#!/usr/bin/env python3
"""
One file only:

1. DOCX -> full_json/*.json
2. optional DeepSeek normalization -> ai_json/*.json

This script never writes the database.
"""

from __future__ import annotations

import argparse
import hashlib
import json
import os
import re
import sqlite3
import sys
import time
import urllib.error
import urllib.request
import zipfile
from datetime import datetime
from pathlib import Path
from typing import Any
from xml.etree import ElementTree as ET

from docx import Document


SCRIPT_DIR = Path(__file__).resolve().parent
SOURCE_DIR = SCRIPT_DIR.parent
FULL_JSON_DIR = SCRIPT_DIR / "full_json"
AI_JSON_DIR = SCRIPT_DIR / "ai_json"
AI_FAILED_DIR = SCRIPT_DIR / "ai_json_failed"
ANNOTATION_DB_DIR = SCRIPT_DIR / "annotation_db"
ANNOTATION_DB_FILE = ANNOTATION_DB_DIR / "annotation_results.db"
DEEPSEEK_URL = "https://api.deepseek.com/chat/completions"
DEFAULT_MODEL = "deepseek-v4-flash"
CURRENT_MAX_TOKENS = 8192

NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
EXTRACTION_SCHEMA = "annotation_docx_full_json_v1"

BLOCK_LABELS = {
    "考據對象": "考据对象",
    "考据对象": "考据对象",
    "反駁對象": "反驳对象",
    "反驳对象": "反驳对象",
    "論點": "论点",
    "论点": "论点",
    "字詞關係": "字词关系",
    "字词关系": "字词关系",
    "引用書證": "引用书证",
    "引用书证": "引用书证",
    "運用術語": "考据用语",
    "考據用語": "考据用语",
    "考据用语": "考据用语",
    "考據過程": "考据过程",
    "考据过程": "考据过程",
    "引用書證的特徵總結": "引用书证的特征总结",
    "引用书证的特征总结": "引用书证的特征总结",
    "整體考據總結": "整体考据总结",
    "整体考据总结": "整体考据总结",
}

METHOD_TAGS = {"校勘", "训释", "声训", "通假", "异体", "异文", "同义互证", "义证", "书证", "形证", "语法证据", "对文散文", "句义解释", "补正"}
RELATION_TYPES = {"", "同义", "通假", "异体", "异文", "声近义同", "误字", "当作", "古今字", "递训", "对文异义", "散文同义"}
EVIDENCE_TYPES = {"书证", "声训", "义证", "形证", "语法证据", "异文"}
STEP_TYPES = {"发疑", "立论", "取证", "释词", "释理", "反驳", "补证", "结论"}

ANNOTATION_SCHEMA = """
PRAGMA foreign_keys = ON;

CREATE TABLE IF NOT EXISTS meta (
    key TEXT PRIMARY KEY,
    value TEXT NOT NULL
);

CREATE TABLE IF NOT EXISTS source_documents (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    source_file_name TEXT NOT NULL UNIQUE,
    source_file_path TEXT,
    full_json_file TEXT,
    ai_json_file TEXT,
    doc_type TEXT,
    source_docx_sha256 TEXT,
    paragraph_text_sha256 TEXT,
    comment_text_sha256 TEXT,
    paragraph_count INTEGER DEFAULT 0,
    comment_count INTEGER DEFAULT 0,
    anchored_comment_count INTEGER DEFAULT 0,
    created_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP,
    updated_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP
);

CREATE TABLE IF NOT EXISTS annotation_cases (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    source_document_id INTEGER REFERENCES source_documents(id),
    case_title TEXT NOT NULL,
    source_work TEXT,
    target_work TEXT,
    target_text TEXT,
    problem TEXT,
    claim TEXT,
    method_tags_json TEXT,
    conclusion TEXT,
    certainty TEXT CHECK(certainty IN ('确定','可疑','待核')) DEFAULT '待核',
    status TEXT CHECK(status IN ('草稿','已校对','已审核')) DEFAULT '草稿',
    raw_case_json TEXT NOT NULL,
    created_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP,
    updated_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP
);

CREATE TABLE IF NOT EXISTS annotation_terms (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    case_id INTEGER NOT NULL REFERENCES annotation_cases(id) ON DELETE CASCADE,
    term TEXT,
    term_type TEXT,
    relation_type TEXT,
    related_term TEXT,
    note TEXT,
    source_paragraph_indexes_json TEXT,
    source_comment_ids_json TEXT,
    raw_term_json TEXT NOT NULL
);

CREATE TABLE IF NOT EXISTS annotation_evidences (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    case_id INTEGER NOT NULL REFERENCES annotation_cases(id) ON DELETE CASCADE,
    evidence_type TEXT CHECK(evidence_type IN ('书证','声训','义证','形证','语法证据','异文')),
    work TEXT,
    quote TEXT,
    role TEXT,
    term TEXT,
    source_paragraph_indexes_json TEXT,
    source_comment_ids_json TEXT,
    raw_evidence_json TEXT NOT NULL
);

CREATE TABLE IF NOT EXISTS annotation_process_steps (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    case_id INTEGER NOT NULL REFERENCES annotation_cases(id) ON DELETE CASCADE,
    step_order INTEGER NOT NULL,
    step_type TEXT,
    text TEXT,
    source_paragraph_indexes_json TEXT,
    source_comment_ids_json TEXT,
    raw_step_json TEXT NOT NULL
);

CREATE INDEX IF NOT EXISTS idx_annotation_cases_source_document_id ON annotation_cases(source_document_id);
CREATE INDEX IF NOT EXISTS idx_annotation_terms_case_id ON annotation_terms(case_id);
CREATE INDEX IF NOT EXISTS idx_annotation_evidences_case_id ON annotation_evidences(case_id);
CREATE INDEX IF NOT EXISTS idx_annotation_process_steps_case_id ON annotation_process_steps(case_id);
"""

SYSTEM_PROMPT = """你是古汉语训诂与校勘标注规范化助手。你的任务是把输入的 DOCX 抽取 JSON 规范化为数据库候选记录。

强制规则：
1. 只能根据输入 JSON 中已有的 paragraphs、comments、annotation_blocks 提取信息。
2. 不得编造书名、引文、字词关系、作者、结论。
3. quote 字段必须复制原文，不得改写。
4. 每个 term、evidence、process_step 必须写 source_paragraph_indexes 或 source_comment_ids。
5. 不确定的内容留空字符串或标为 待核，不要猜。
6. 只输出 JSON，不输出 Markdown，不输出解释。"""

USER_PROMPT_PREFIX = """请把下面的 DOCX 抽取 JSON 规范化为考据案例 JSON。

输出必须是 JSON，根对象格式如下：
{
  "cases": [
    {
      "case_title": "",
      "source_work": "",
      "target_work": "",
      "target_text": "",
      "problem": "",
      "claim": "",
      "method_tags": [],
      "terms": [
        {
          "term": "",
          "term_type": "",
          "relation_type": "",
          "related_term": "",
          "note": "",
          "source_paragraph_indexes": [],
          "source_comment_ids": []
        }
      ],
      "evidences": [
        {
          "evidence_type": "",
          "work": "",
          "quote": "",
          "role": "",
          "term": "",
          "source_paragraph_indexes": [],
          "source_comment_ids": []
        }
      ],
      "process_steps": [
        {
          "step_type": "",
          "text": "",
          "source_paragraph_indexes": [],
          "source_comment_ids": []
        }
      ],
      "conclusion": "",
      "certainty": "待核",
      "status": "草稿"
    }
  ]
}

枚举限制：
- method_tags: 校勘, 训释, 声训, 通假, 异体, 异文, 同义互证, 义证, 书证, 形证, 语法证据, 对文散文, 句义解释, 补正
- relation_type: 同义, 通假, 异体, 异文, 声近义同, 误字, 当作, 古今字, 递训, 对文异义, 散文同义
- evidence_type: 书证, 声训, 义证, 形证, 语法证据, 异文
- step_type: 发疑, 立论, 取证, 释词, 释理, 反驳, 补证, 结论
- certainty: 确定, 可疑, 待核
- status: 草稿, 已校对, 已审核

输入 JSON：
"""


def load_env_file() -> None:
    """Load optional local .env and prefer it over stale shell variables."""
    env_file = SCRIPT_DIR / ".env"
    if not env_file.exists():
        return
    for raw_line in env_file.read_text(encoding="utf-8").splitlines():
        line = raw_line.strip()
        if not line or line.startswith("#") or "=" not in line:
            continue
        key, value = line.split("=", 1)
        key = key.strip()
        value = value.strip().strip('"').strip("'")
        if key:
            os.environ[key] = value


def api_key_fingerprint() -> str:
    api_key = os.environ.get("DEEPSEEK_API_KEY", "").strip()
    if not api_key:
        return "missing"
    digest = hashlib.sha256(api_key.encode("utf-8")).hexdigest()[:12]
    tail = api_key[-4:] if len(api_key) >= 4 else api_key
    return f"len={len(api_key)} sha12={digest} tail4={tail}"


def init_annotation_db() -> None:
    ANNOTATION_DB_DIR.mkdir(exist_ok=True)
    with sqlite3.connect(str(ANNOTATION_DB_FILE), isolation_level=None) as conn:
        conn.execute("PRAGMA journal_mode=WAL")
        conn.execute("PRAGMA synchronous=NORMAL")
        conn.executescript(ANNOTATION_SCHEMA)
        conn.execute("INSERT OR REPLACE INTO meta(key, value) VALUES(?, ?)", ("schema_version", "annotation_db_v1"))
        conn.execute("INSERT OR REPLACE INTO meta(key, value) VALUES(?, ?)", ("purpose", "人工标注与 AI 整理结果暂存库，不与 02-数据库 主库混用"))


def qname(tag: str) -> str:
    return f"{{{NS['w']}}}{tag}"


def get_attr(element: ET.Element, name: str) -> str | None:
    return element.attrib.get(qname(name))


def text_of(element: ET.Element | None) -> str:
    if element is None:
        return ""
    return "".join(node.text or "" for node in element.iter(qname("t")))


def sha256_bytes(value: bytes) -> str:
    return hashlib.sha256(value).hexdigest()


def sha256_text(value: str) -> str:
    return sha256_bytes(value.encode("utf-8"))


def parse_comments(docx_path: Path) -> dict[str, dict[str, Any]]:
    with zipfile.ZipFile(docx_path) as archive:
        if "word/comments.xml" not in archive.namelist():
            return {}
        root = ET.fromstring(archive.read("word/comments.xml"))

    comments: dict[str, dict[str, Any]] = {}
    for comment in root.findall("w:comment", NS):
        comment_id = get_attr(comment, "id")
        if comment_id is None:
            continue
        comments[comment_id] = {
            "id": comment_id,
            "author": get_attr(comment, "author") or "",
            "date": get_attr(comment, "date") or "",
            "text": text_of(comment).strip(),
            "anchor_text": "",
            "paragraph_indexes": [],
        }
    return comments


def parse_comment_anchors(docx_path: Path, comments: dict[str, dict[str, Any]]) -> None:
    if not comments:
        return

    with zipfile.ZipFile(docx_path) as archive:
        root = ET.fromstring(archive.read("word/document.xml"))

    active: dict[str, list[str]] = {}
    paragraph_index = 0
    for paragraph in root.iter(qname("p")):
        paragraph_index += 1
        ids_in_paragraph: set[str] = set()
        for node in paragraph.iter():
            if node.tag == qname("commentRangeStart"):
                comment_id = get_attr(node, "id")
                if comment_id in comments:
                    active[comment_id] = []
                    ids_in_paragraph.add(comment_id)
            elif node.tag == qname("t"):
                for parts in active.values():
                    parts.append(node.text or "")
            elif node.tag == qname("commentRangeEnd"):
                comment_id = get_attr(node, "id")
                if comment_id in active and comment_id in comments:
                    comments[comment_id]["anchor_text"] = "".join(active.pop(comment_id)).strip()
                    ids_in_paragraph.add(comment_id)
            elif node.tag == qname("commentReference"):
                comment_id = get_attr(node, "id")
                if comment_id in comments:
                    ids_in_paragraph.add(comment_id)

        for comment_id in ids_in_paragraph:
            comments[comment_id]["paragraph_indexes"].append(paragraph_index)


def paragraph_comment_map(docx_path: Path) -> dict[int, list[str]]:
    with zipfile.ZipFile(docx_path) as archive:
        root = ET.fromstring(archive.read("word/document.xml"))

    mapping: dict[int, list[str]] = {}
    paragraph_index = 0
    for paragraph in root.iter(qname("p")):
        paragraph_index += 1
        ids: list[str] = []
        for tag_name in ("commentRangeStart", "commentRangeEnd", "commentReference"):
            for node in paragraph.iter(qname(tag_name)):
                comment_id = get_attr(node, "id")
                if comment_id is not None and comment_id not in ids:
                    ids.append(comment_id)
        if ids:
            mapping[paragraph_index] = ids
    return mapping


def extract_run(run: Any) -> dict[str, Any]:
    return {
        "text": run.text,
        "bold": bool(run.bold) if run.bold is not None else None,
        "italic": bool(run.italic) if run.italic is not None else None,
        "underline": bool(run.underline) if run.underline is not None else None,
        "style": run.style.name if run.style else "",
        "font_name": run.font.name,
    }


def extract_paragraphs(docx_path: Path, comment_map: dict[int, list[str]]) -> list[dict[str, Any]]:
    document = Document(str(docx_path))
    paragraphs = []
    for index, paragraph in enumerate(document.paragraphs, start=1):
        paragraphs.append(
            {
                "index": index,
                "style": paragraph.style.name if paragraph.style else "",
                "text": paragraph.text,
                "is_empty": not bool(paragraph.text.strip()),
                "runs": [extract_run(run) for run in paragraph.runs],
                "comment_ids": comment_map.get(index, []),
            }
        )
    return paragraphs


def extract_tables(docx_path: Path) -> list[dict[str, Any]]:
    document = Document(str(docx_path))
    tables = []
    for table_index, table in enumerate(document.tables, start=1):
        rows = []
        for row in table.rows:
            rows.append(
                [
                    {"text": cell.text, "paragraphs": [paragraph.text for paragraph in cell.paragraphs]}
                    for cell in row.cells
                ]
            )
        tables.append({"index": table_index, "rows": rows})
    return tables


def extract_notes(docx_path: Path) -> dict[str, list[dict[str, str]]]:
    result = {"footnotes": [], "endnotes": []}
    files = {"footnotes": "word/footnotes.xml", "endnotes": "word/endnotes.xml"}
    with zipfile.ZipFile(docx_path) as archive:
        names = set(archive.namelist())
        for note_type, note_file in files.items():
            if note_file not in names:
                continue
            root = ET.fromstring(archive.read(note_file))
            item_tag = "footnote" if note_type == "footnotes" else "endnote"
            for note in root.findall(f"w:{item_tag}", NS):
                note_text = text_of(note).strip()
                if note_text:
                    result[note_type].append({"id": get_attr(note, "id") or "", "text": note_text})
    return result


def package_inventory(docx_path: Path) -> list[dict[str, Any]]:
    with zipfile.ZipFile(docx_path) as archive:
        return [
            {
                "name": info.filename,
                "size": info.file_size,
                "compressed_size": info.compress_size,
                "crc": f"{info.CRC:08x}",
            }
            for info in sorted(archive.infolist(), key=lambda item: item.filename)
        ]


def normalize_label(text: str) -> str | None:
    return BLOCK_LABELS.get(text.strip().strip("：:"))


def build_annotation_blocks(paragraphs: list[dict[str, Any]]) -> list[dict[str, Any]]:
    blocks = []
    current: dict[str, Any] | None = None
    for paragraph in paragraphs:
        if paragraph["is_empty"]:
            continue
        label = normalize_label(paragraph["text"])
        if label:
            if current:
                current["end_paragraph"] = paragraph["index"] - 1
                blocks.append(current)
            current = {"label": label, "start_paragraph": paragraph["index"], "end_paragraph": paragraph["index"], "items": []}
            continue
        if current:
            current["items"].append(paragraph["text"])
            current["end_paragraph"] = paragraph["index"]
    if current:
        blocks.append(current)
    return blocks


def infer_doc_type(filename: str, paragraphs: list[dict[str, Any]], comments: list[dict[str, Any]]) -> str:
    first_text = "\n".join(item["text"] for item in paragraphs[:30])
    if comments and "经传释词" in filename:
        return "word_comments_first"
    if "读书杂志" in filename:
        return "case_blocks"
    if "广雅疏证" in filename:
        return "term_group_blocks"
    if re.search(r"考據對象|考据对象|引用書證|引用书证", first_text):
        return "case_blocks"
    return "plain_docx"


def convert_one(docx_path: Path) -> dict[str, Any]:
    comments_map = parse_comments(docx_path)
    parse_comment_anchors(docx_path, comments_map)
    paragraphs = extract_paragraphs(docx_path, paragraph_comment_map(docx_path))
    tables = extract_tables(docx_path)
    notes = extract_notes(docx_path)
    comments = list(comments_map.values())
    paragraph_text = "\n".join(item["text"] for item in paragraphs)
    comment_text = "\n".join(item["text"] for item in comments)
    stat = docx_path.stat()

    return {
        "schema_version": EXTRACTION_SCHEMA,
        "source_file": {
            "name": docx_path.name,
            "path": str(docx_path.relative_to(SOURCE_DIR.parent)),
            "size_bytes": stat.st_size,
            "modified_time": datetime.fromtimestamp(stat.st_mtime).isoformat(timespec="seconds"),
        },
        "document_stats": {
            "paragraph_count": len(paragraphs),
            "nonempty_paragraph_count": sum(1 for item in paragraphs if not item["is_empty"]),
            "table_count": len(tables),
            "comment_count": len(comments),
            "anchored_comment_count": sum(1 for item in comments if item.get("anchor_text")),
            "annotation_block_count": len(build_annotation_blocks(paragraphs)),
            "footnote_count": len(notes["footnotes"]),
            "endnote_count": len(notes["endnotes"]),
            "paragraph_text_length": sum(len(item["text"]) for item in paragraphs),
            "comment_text_length": sum(len(item["text"]) for item in comments),
        },
        "doc_type": infer_doc_type(docx_path.name, paragraphs, comments),
        "paragraphs": paragraphs,
        "tables": tables,
        "comments": comments,
        "notes": notes,
        "annotation_blocks": build_annotation_blocks(paragraphs),
        "package_inventory": package_inventory(docx_path),
        "checksums": {
            "source_docx_sha256": sha256_bytes(docx_path.read_bytes()),
            "paragraph_text_sha256": sha256_text(paragraph_text),
            "comment_text_sha256": sha256_text(comment_text),
        },
        "controlled_vocabularies": {
            "method_tags": sorted(METHOD_TAGS),
            "relation_types": sorted(RELATION_TYPES),
            "evidence_types": sorted(EVIDENCE_TYPES),
            "step_types": sorted(STEP_TYPES),
            "certainty": ["确定", "可疑", "待核"],
            "status": ["草稿", "已校对", "已审核"],
        },
    }


def convert_docx_files(only: str | None = None) -> list[Path]:
    FULL_JSON_DIR.mkdir(exist_ok=True)
    docx_files = sorted(SOURCE_DIR.glob("*.docx"), key=lambda item: item.name)
    if only:
        docx_files = [path for path in docx_files if only in path.name or only in path.stem]
    if not docx_files:
        raise RuntimeError("No matching DOCX files found.")

    outputs = []
    report = []
    for docx_path in docx_files:
        payload = convert_one(docx_path)
        output_path = FULL_JSON_DIR / f"{docx_path.stem}.json"
        output_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
        outputs.append(output_path)
        report.append({"source_file": docx_path.name, "json_file": output_path.name, "stats": payload["document_stats"], "checksums": payload["checksums"]})
        stats = payload["document_stats"]
        print(f"full_json: {output_path.name} paragraphs={stats['paragraph_count']} comments={stats['comment_count']} anchors={stats['anchored_comment_count']}")

    (FULL_JSON_DIR / "_report.json").write_text(json.dumps({"generated_at": datetime.now().isoformat(timespec="seconds"), "documents": report}, ensure_ascii=False, indent=2), encoding="utf-8")
    return sorted(outputs, key=lambda path: path.stat().st_size)


def validate_full_json(path: Path) -> None:
    data = json.loads(path.read_text(encoding="utf-8"))
    paragraphs = data["paragraphs"]
    comments = data["comments"]
    stats = data["document_stats"]
    checksums = data["checksums"]
    paragraph_text = "\n".join(item["text"] for item in paragraphs)
    comment_text = "\n".join(item["text"] for item in comments)

    checks = {
        "paragraph_count": len(paragraphs),
        "nonempty_paragraph_count": sum(1 for item in paragraphs if not item["is_empty"]),
        "comment_count": len(comments),
        "anchored_comment_count": sum(1 for item in comments if item.get("anchor_text")),
        "paragraph_text_length": sum(len(item["text"]) for item in paragraphs),
        "comment_text_length": sum(len(item["text"]) for item in comments),
    }
    for key, value in checks.items():
        if stats.get(key) != value:
            raise RuntimeError(f"{path.name}: {key} mismatch, stats={stats.get(key)!r}, actual={value!r}")
    if checksums["paragraph_text_sha256"] != sha256_text(paragraph_text):
        raise RuntimeError(f"{path.name}: paragraph checksum mismatch")
    if checksums["comment_text_sha256"] != sha256_text(comment_text):
        raise RuntimeError(f"{path.name}: comment checksum mismatch")


def slim_for_ai(data: dict[str, Any]) -> dict[str, Any]:
    return {
        "source_file": data["source_file"],
        "document_stats": data["document_stats"],
        "doc_type": data["doc_type"],
        "paragraphs": [
            {"index": item["index"], "text": item["text"], "comment_ids": item.get("comment_ids", [])}
            for item in data["paragraphs"]
            if not item["is_empty"] and item["text"].strip()
        ],
        "comments": [
            {"id": item["id"], "text": item.get("text", ""), "anchor_text": item.get("anchor_text", ""), "paragraph_indexes": item.get("paragraph_indexes", [])}
            for item in data["comments"]
        ],
        "annotation_blocks": data["annotation_blocks"],
        "controlled_vocabularies": data["controlled_vocabularies"],
    }


def call_deepseek(full_json_path: Path, model: str, timeout: int, retries: int) -> dict[str, Any]:
    api_key = os.environ.get("DEEPSEEK_API_KEY", "").strip()
    if not api_key:
        raise RuntimeError('Missing DEEPSEEK_API_KEY. PowerShell: $env:DEEPSEEK_API_KEY = "sk-..."')

    data = json.loads(full_json_path.read_text(encoding="utf-8"))
    source_payload = json.dumps(slim_for_ai(data), ensure_ascii=False)
    request_payload = {
        "model": model,
        "response_format": {"type": "json_object"},
        "temperature": 0,
        "max_tokens": CURRENT_MAX_TOKENS,
        "messages": [
            {"role": "system", "content": SYSTEM_PROMPT},
            {"role": "user", "content": USER_PROMPT_PREFIX + "\n```json\n" + source_payload + "\n```"},
        ],
    }
    request = urllib.request.Request(
        DEEPSEEK_URL,
        data=json.dumps(request_payload, ensure_ascii=False).encode("utf-8"),
        headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
        method="POST",
    )

    last_error: Exception | None = None
    for attempt in range(1, retries + 2):
        try:
            with urllib.request.urlopen(request, timeout=timeout) as response:
                response_data = json.loads(response.read().decode("utf-8"))
            content = response_data["choices"][0]["message"].get("content", "")
            if not content.strip():
                raise RuntimeError("empty API content")
            try:
                return json.loads(content)
            except json.JSONDecodeError as error:
                AI_FAILED_DIR.mkdir(exist_ok=True)
                failed_path = AI_FAILED_DIR / f"{full_json_path.stem}.raw.txt"
                failed_path.write_text(content, encoding="utf-8")
                raise RuntimeError(f"invalid JSON content saved to {failed_path}: {error}") from error
        except urllib.error.HTTPError as error:
            if error.code == 401:
                raise RuntimeError(
                    "DeepSeek authentication failed: HTTP 401. "
                    "Check DEEPSEEK_API_KEY in json/.env or your account permissions."
                ) from error
            last_error = error
            if attempt > retries:
                break
            wait = min(2 * attempt, 10)
            print(f"API failed for {full_json_path.name}: {error}; retrying in {wait}s")
            time.sleep(wait)
        except (urllib.error.URLError, TimeoutError, json.JSONDecodeError, KeyError, RuntimeError) as error:
            last_error = error
            if attempt > retries:
                break
            wait = min(2 * attempt, 10)
            print(f"API failed for {full_json_path.name}: {error}; retrying in {wait}s")
            time.sleep(wait)
    raise RuntimeError(f"DeepSeek failed for {full_json_path.name}: {last_error}")


def check_api_key(model: str, timeout: int) -> None:
    api_key = os.environ.get("DEEPSEEK_API_KEY", "").strip()
    if not api_key:
        raise RuntimeError('Missing DEEPSEEK_API_KEY. Put it in json/.env as DEEPSEEK_API_KEY=sk-...')

    payload = {
        "model": model,
        "messages": [{"role": "user", "content": "ping"}],
        "max_tokens": 1,
        "temperature": 0,
    }
    request = urllib.request.Request(
        DEEPSEEK_URL,
        data=json.dumps(payload, ensure_ascii=False).encode("utf-8"),
        headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
        method="POST",
    )
    try:
        with urllib.request.urlopen(request, timeout=timeout) as response:
            json.loads(response.read().decode("utf-8"))
        print(f"DeepSeek API key OK ({api_key_fingerprint()})")
    except urllib.error.HTTPError as error:
        if error.code == 401:
            raise RuntimeError(f"DeepSeek API key rejected: HTTP 401 ({api_key_fingerprint()})") from error
        raise


def validate_ai_json(data: dict[str, Any], name: str) -> None:
    cases = data.get("cases")
    if not isinstance(cases, list):
        raise RuntimeError(f"{name}: root must contain cases array")
    for case_index, case in enumerate(cases, start=1):
        for tag in case.get("method_tags", []):
            if tag not in METHOD_TAGS:
                raise RuntimeError(f"{name} case {case_index}: invalid method tag {tag}")
        for term in case.get("terms", []):
            if term.get("relation_type", "") not in RELATION_TYPES:
                raise RuntimeError(f"{name} case {case_index}: invalid relation_type {term.get('relation_type')}")
        for evidence in case.get("evidences", []):
            if evidence.get("evidence_type") not in EVIDENCE_TYPES:
                raise RuntimeError(f"{name} case {case_index}: invalid evidence_type {evidence.get('evidence_type')}")
            if not evidence.get("quote", "").strip():
                raise RuntimeError(f"{name} case {case_index}: empty evidence quote")
        for step in case.get("process_steps", []):
            if step.get("step_type") not in STEP_TYPES:
                raise RuntimeError(f"{name} case {case_index}: invalid step_type {step.get('step_type')}")


def add_ingestion_metadata(data: dict[str, Any], imported: bool = False) -> dict[str, Any]:
    data.setdefault(
        "database_ingestion",
        {
            "imported": imported,
            "database": str(ANNOTATION_DB_FILE),
            "imported_at": None,
        },
    )
    data["database_ingestion"]["imported"] = imported
    data["database_ingestion"]["database"] = str(ANNOTATION_DB_FILE)
    if not imported:
        data["database_ingestion"]["imported_at"] = None

    for case in data.get("cases", []):
        case.setdefault(
            "database_ingestion",
            {
                "imported": imported,
                "annotation_case_id": None,
                "imported_at": None,
            },
        )
        case["database_ingestion"]["imported"] = imported
        if not imported:
            case["database_ingestion"]["annotation_case_id"] = None
            case["database_ingestion"]["imported_at"] = None
    return data


def normalize_with_ai(full_json_paths: list[Path], model: str, timeout: int, retries: int) -> list[Path]:
    AI_JSON_DIR.mkdir(exist_ok=True)
    outputs = []
    for full_json_path in full_json_paths:
        print(f"DeepSeek: {full_json_path.name}")
        payload = call_deepseek(full_json_path, model=model, timeout=timeout, retries=retries)
        validate_ai_json(payload, full_json_path.name)
        payload = add_ingestion_metadata(payload, imported=False)
        output_path = AI_JSON_DIR / full_json_path.name
        output_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
        outputs.append(output_path)
        print(f"ai_json: {output_path.name}")
    return outputs


def existing_full_json(only: str | None = None) -> list[Path]:
    paths = sorted(
        (path for path in FULL_JSON_DIR.glob("*.json") if path.name != "_report.json"),
        key=lambda path: path.stat().st_size,
    )
    if only:
        paths = [path for path in paths if only in path.name or only in path.stem]
    if not paths:
        raise RuntimeError("No full_json files found. Run without --skip-convert first.")
    return paths


def existing_ai_json(only: str | None = None) -> list[Path]:
    paths = sorted(path for path in AI_JSON_DIR.glob("*.json"))
    if only:
        paths = [path for path in paths if only in path.name or only in path.stem]
    if not paths:
        raise RuntimeError("No ai_json files found. Run with --api first.")
    return paths


def json_dumps(value: Any) -> str:
    return json.dumps(value, ensure_ascii=False)


def import_ai_json_files(ai_json_paths: list[Path]) -> None:
    init_annotation_db()
    imported_at = datetime.now().isoformat(timespec="seconds")

    with sqlite3.connect(str(ANNOTATION_DB_FILE)) as conn:
        conn.execute("PRAGMA foreign_keys = ON")
        for ai_json_path in ai_json_paths:
            data = json.loads(ai_json_path.read_text(encoding="utf-8"))
            validate_ai_json(data, ai_json_path.name)

            full_json_path = FULL_JSON_DIR / ai_json_path.name
            if not full_json_path.exists():
                raise RuntimeError(f"Missing matching full_json file for {ai_json_path.name}")
            full_data = json.loads(full_json_path.read_text(encoding="utf-8"))
            source = full_data["source_file"]
            stats = full_data["document_stats"]
            checksums = full_data["checksums"]

            conn.execute(
                """
                INSERT INTO source_documents(
                    source_file_name, source_file_path, full_json_file, ai_json_file,
                    doc_type, source_docx_sha256, paragraph_text_sha256, comment_text_sha256,
                    paragraph_count, comment_count, anchored_comment_count, updated_at
                )
                VALUES(?,?,?,?,?,?,?,?,?,?,?,CURRENT_TIMESTAMP)
                ON CONFLICT(source_file_name) DO UPDATE SET
                    source_file_path=excluded.source_file_path,
                    full_json_file=excluded.full_json_file,
                    ai_json_file=excluded.ai_json_file,
                    doc_type=excluded.doc_type,
                    source_docx_sha256=excluded.source_docx_sha256,
                    paragraph_text_sha256=excluded.paragraph_text_sha256,
                    comment_text_sha256=excluded.comment_text_sha256,
                    paragraph_count=excluded.paragraph_count,
                    comment_count=excluded.comment_count,
                    anchored_comment_count=excluded.anchored_comment_count,
                    updated_at=CURRENT_TIMESTAMP
                """,
                (
                    source["name"],
                    source.get("path", ""),
                    str(full_json_path.relative_to(SCRIPT_DIR)),
                    str(ai_json_path.relative_to(SCRIPT_DIR)),
                    full_data.get("doc_type", ""),
                    checksums.get("source_docx_sha256", ""),
                    checksums.get("paragraph_text_sha256", ""),
                    checksums.get("comment_text_sha256", ""),
                    stats.get("paragraph_count", 0),
                    stats.get("comment_count", 0),
                    stats.get("anchored_comment_count", 0),
                ),
            )
            source_document_id = conn.execute(
                "SELECT id FROM source_documents WHERE source_file_name=?",
                (source["name"],),
            ).fetchone()[0]

            conn.execute("DELETE FROM annotation_cases WHERE source_document_id=?", (source_document_id,))

            case_ids: list[int] = []
            for case in data.get("cases", []):
                cursor = conn.execute(
                    """
                    INSERT INTO annotation_cases(
                        source_document_id, case_title, source_work, target_work, target_text,
                        problem, claim, method_tags_json, conclusion, certainty, status, raw_case_json
                    )
                    VALUES(?,?,?,?,?,?,?,?,?,?,?,?)
                    """,
                    (
                        source_document_id,
                        case.get("case_title") or "未命名标注案例",
                        case.get("source_work", ""),
                        case.get("target_work", ""),
                        case.get("target_text", ""),
                        case.get("problem", ""),
                        case.get("claim", ""),
                        json_dumps(case.get("method_tags", [])),
                        case.get("conclusion", ""),
                        case.get("certainty") or "待核",
                        case.get("status") or "草稿",
                        json_dumps(case),
                    ),
                )
                case_id = int(cursor.lastrowid)
                case_ids.append(case_id)

                for term in case.get("terms", []):
                    conn.execute(
                        """
                        INSERT INTO annotation_terms(
                            case_id, term, term_type, relation_type, related_term, note,
                            source_paragraph_indexes_json, source_comment_ids_json, raw_term_json
                        )
                        VALUES(?,?,?,?,?,?,?,?,?)
                        """,
                        (
                            case_id,
                            term.get("term", ""),
                            term.get("term_type", ""),
                            term.get("relation_type", ""),
                            term.get("related_term", ""),
                            term.get("note", ""),
                            json_dumps(term.get("source_paragraph_indexes", [])),
                            json_dumps(term.get("source_comment_ids", [])),
                            json_dumps(term),
                        ),
                    )

                for evidence in case.get("evidences", []):
                    conn.execute(
                        """
                        INSERT INTO annotation_evidences(
                            case_id, evidence_type, work, quote, role, term,
                            source_paragraph_indexes_json, source_comment_ids_json, raw_evidence_json
                        )
                        VALUES(?,?,?,?,?,?,?,?,?)
                        """,
                        (
                            case_id,
                            evidence.get("evidence_type", "书证"),
                            evidence.get("work", ""),
                            evidence.get("quote", ""),
                            evidence.get("role", ""),
                            evidence.get("term", ""),
                            json_dumps(evidence.get("source_paragraph_indexes", [])),
                            json_dumps(evidence.get("source_comment_ids", [])),
                            json_dumps(evidence),
                        ),
                    )

                for step_order, step in enumerate(case.get("process_steps", []), start=1):
                    conn.execute(
                        """
                        INSERT INTO annotation_process_steps(
                            case_id, step_order, step_type, text,
                            source_paragraph_indexes_json, source_comment_ids_json, raw_step_json
                        )
                        VALUES(?,?,?,?,?,?,?)
                        """,
                        (
                            case_id,
                            step_order,
                            step.get("step_type", ""),
                            step.get("text", ""),
                            json_dumps(step.get("source_paragraph_indexes", [])),
                            json_dumps(step.get("source_comment_ids", [])),
                            json_dumps(step),
                        ),
                    )

            data = add_ingestion_metadata(data, imported=True)
            data["database_ingestion"]["imported_at"] = imported_at
            for case, case_id in zip(data.get("cases", []), case_ids):
                case["database_ingestion"]["imported"] = True
                case["database_ingestion"]["annotation_case_id"] = case_id
                case["database_ingestion"]["imported_at"] = imported_at
            ai_json_path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
            print(f"imported: {ai_json_path.name} cases={len(case_ids)}")

        conn.commit()


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="DOCX -> full_json, optional DeepSeek -> ai_json. Does not touch database.")
    parser.add_argument("--api", action="store_true", help="Call DeepSeek and write ai_json/*.json.")
    parser.add_argument("--check-api-key", action="store_true", help="Send a tiny DeepSeek request to verify the configured key.")
    parser.add_argument("--import-ai", action="store_true", help="Import ai_json/*.json into the separate annotation database.")
    parser.add_argument("--init-db", action="store_true", help="Only initialize the separate annotation SQLite database.")
    parser.add_argument("--skip-convert", action="store_true", help="Use existing full_json/*.json.")
    parser.add_argument("--only", help="Only process files whose name contains this text.")
    parser.add_argument("--model", default=DEFAULT_MODEL, help=f"DeepSeek model. Default: {DEFAULT_MODEL}")
    parser.add_argument("--max-tokens", type=int, default=8192, help="DeepSeek max output tokens. Default: 8192")
    parser.add_argument("--timeout", type=int, default=180, help="API timeout seconds.")
    parser.add_argument("--retries", type=int, default=1, help="API retries.")
    return parser.parse_args()


def main() -> int:
    global CURRENT_MAX_TOKENS
    args = parse_args()
    CURRENT_MAX_TOKENS = args.max_tokens
    load_env_file()
    if args.check_api_key:
        check_api_key(model=args.model, timeout=args.timeout)
        return 0
    if args.init_db:
        init_annotation_db()
        print(f"annotation db ready: {ANNOTATION_DB_FILE}")
        return 0
    if not ANNOTATION_DB_FILE.exists():
        init_annotation_db()
        print(f"annotation db initialized: {ANNOTATION_DB_FILE}")
    full_json_paths = existing_full_json(args.only) if args.skip_convert else convert_docx_files(args.only)
    for path in full_json_paths:
        validate_full_json(path)
    print("full_json validation OK")
    if args.api:
        ai_json_paths = normalize_with_ai(full_json_paths, model=args.model, timeout=args.timeout, retries=args.retries)
        print("ai_json validation OK")
        if args.import_ai:
            import_ai_json_files(ai_json_paths)
            print(f"annotation db import OK: {ANNOTATION_DB_FILE}")
    elif args.import_ai:
        ai_json_paths = existing_ai_json(args.only)
        import_ai_json_files(ai_json_paths)
        print(f"annotation db import OK: {ANNOTATION_DB_FILE}")
    else:
        print("Done. Add --api to generate ai_json with DeepSeek.")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
