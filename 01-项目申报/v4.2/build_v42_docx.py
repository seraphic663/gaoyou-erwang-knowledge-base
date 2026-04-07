from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


BASE_DIR = Path(__file__).resolve().parent
SOURCE_MD = BASE_DIR / "立项申报v4.2.md"
TEMPLATE_DOCX = BASE_DIR / "立项申报v4.1（格式调整版）-批注落实副本.docx"
OUTPUT_DOCX = BASE_DIR / "立项申报v4.2.docx"


def parse_markdown(text: str) -> list[dict]:
    lines = text.splitlines()
    blocks: list[dict] = []
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()
        if not stripped:
            i += 1
            continue

        if stripped.startswith("|") and "|" in stripped[1:]:
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            blocks.append({"type": "table", "lines": table_lines})
            continue

        heading_match = re.match(r"^(#{1,4})\s+(.*)$", stripped)
        if heading_match:
            blocks.append(
                {
                    "type": "heading",
                    "level": len(heading_match.group(1)),
                    "text": heading_match.group(2).strip(),
                }
            )
            i += 1
            continue

        bullet_match = re.match(r"^-\s+(.*)$", stripped)
        if bullet_match:
            items = []
            while i < len(lines):
                item_line = lines[i].strip()
                m = re.match(r"^-\s+(.*)$", item_line)
                if not m:
                    break
                items.append(m.group(1).strip())
                i += 1
            blocks.append({"type": "bullet_list", "items": items})
            continue

        num_match = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if num_match:
            items = []
            while i < len(lines):
                item_line = lines[i].strip()
                m = re.match(r"^(\d+)\.\s+(.*)$", item_line)
                if not m:
                    break
                items.append(m.group(0).strip())
                i += 1
            blocks.append({"type": "number_list", "items": items})
            continue

        para_lines = [stripped]
        i += 1
        while i < len(lines):
            nxt = lines[i].strip()
            if not nxt:
                i += 1
                break
            if nxt.startswith("|") or nxt.startswith("#") or nxt.startswith("- "):
                break
            if re.match(r"^\d+\.\s+", nxt):
                break
            para_lines.append(nxt)
            i += 1
        blocks.append({"type": "paragraph", "text": "".join(para_lines)})
    return blocks


def clear_document(doc: Document) -> None:
    body = doc._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def build_docx() -> None:
    markdown = SOURCE_MD.read_text(encoding="utf-8")
    blocks = parse_markdown(markdown)

    template = Document(str(TEMPLATE_DOCX))
    table_style = template.tables[0].style if template.tables else None
    clear_document(template)

    current_h3 = ""
    current_h2 = ""

    for block in blocks:
        btype = block["type"]
        if btype == "heading":
            level = block["level"]
            text = block["text"]
            style = {
                1: "Heading 1",
                2: "Heading 2",
                3: "Heading 3",
                4: "Heading 4",
            }[level]
            template.add_paragraph(text, style=style)
            if level == 2:
                current_h2 = text
                current_h3 = ""
            elif level == 3:
                current_h3 = text
            continue

        if btype == "paragraph":
            text = block["text"].replace("**", "")
            style = "Normal"
            if text.startswith("【说明】"):
                style = "First Paragraph"
            elif current_h2 == "六、项目经费预算" and text.startswith("本项目坚持经费使用"):
                style = "Body Text"
            elif current_h2 == "七、项目预期研究成果" and text.startswith("项目拟形成以下成果："):
                style = "Normal"
            template.add_paragraph(text, style=style)
            continue

        if btype == "bullet_list":
            for item in block["items"]:
                item = item.replace("**", "")
                style = "List Paragraph" if current_h3 == "5. 主要参考文献" else "Normal"
                template.add_paragraph(item, style=style)
            continue

        if btype == "number_list":
            for item in block["items"]:
                item = item.replace("**", "")
                style = "List Paragraph" if current_h3 == "5. 主要参考文献" else "Normal"
                template.add_paragraph(item, style=style)
            continue

        if btype == "table":
            rows = []
            for line in block["lines"]:
                cells = [c.strip() for c in line.strip("|").split("|")]
                if set("".join(cells)) == {"-"}:
                    continue
                rows.append(cells)
            if not rows:
                continue
            table = template.add_table(rows=len(rows), cols=len(rows[0]))
            if table_style is not None:
                table.style = table_style
            for r_idx, row in enumerate(rows):
                for c_idx, cell in enumerate(row):
                    table.cell(r_idx, c_idx).text = cell
            continue

    template.save(str(OUTPUT_DOCX))


if __name__ == "__main__":
    build_docx()
